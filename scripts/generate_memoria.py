"""Genera la Memoria TFG Entrega 1 en formato Word."""

import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_PATH = os.path.join(ROOT, "docs", "Memoria_TFG_Entrega1.docx")

# ── Colores ────────────────────────────────────────────────────────────────────
BLUE_DARK   = RGBColor(0x1F, 0x4E, 0x8A)   # títulos capítulo
BLUE_MED    = RGBColor(0x2E, 0x75, 0xB6)   # títulos sección
BLUE_LIGHT  = RGBColor(0xD6, 0xE4, 0xF0)   # filas alternas tabla
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
GRAY_BG     = RGBColor(0xF2, 0xF2, 0xF2)   # fondo código

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, color_hex: str):
    """Pone color de fondo a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_heading(doc: Document, text: str, level: int):
    """Añade un título con el estilo de la memoria."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = BLUE_DARK
        p.paragraph_format.page_break_before = True
    else:
        run.font.size = Pt(13)
        run.font.color.rgb = BLUE_MED

def add_body(doc: Document, text: str):
    """Añade párrafo de cuerpo."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0.5)
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
    return p

def add_code_block(doc: Document, code: str):
    """Añade un bloque de código con estilo Courier New."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    # Fondo gris
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

def add_table(doc: Document, headers: list, rows: list):
    """Añade tabla con encabezado azul oscuro y filas alternas."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Encabezado
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, '1F4E8A')
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(10)
    # Filas
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = 'FFFFFF' if r_idx % 2 == 0 else 'D6E4F0'
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()

def add_page_number(doc: Document):
    """Añade numeración de página al pie centrada."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_header(doc: Document, title: str):
    """Añade encabezado con el título del TFG."""
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = title
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = BLUE_MED
        run.font.italic = True

# ══════════════════════════════════════════════════════════════════════════════
#  GENERACIÓN DEL DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════

def build_document():
    doc = Document()

    # ── Márgenes ──────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Estilo base del cuerpo ────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = Pt(18)  # interlineado 1.5

    add_header(doc, "Asistente Universitario Inteligente basado en RAG y Orquestación de Agentes")
    add_page_number(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PORTADA
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("Asistente Universitario Inteligente\nbasado en RAG y Orquestación de Agentes\ncon LangChain y LangGraph")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = BLUE_DARK

    doc.add_paragraph()
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run("Trabajo de Fin de Grado — Memoria Parcial (Entrega 1)")
    r2.font.size = Pt(14)
    r2.font.color.rgb = BLUE_MED

    doc.add_paragraph()
    doc.add_paragraph()

    for label, value in [
        ("Autor:", "Félix García"),
        ("Tutor:", "[Por definir]"),
        ("Titulación:", "Ingeniería Informática"),
        ("Fecha:", "Abril 2025"),
        ("Institución:", "[Universidad]"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(f"{label}  ")
        rl.font.bold = True
        rl.font.size = Pt(12)
        rv = p.add_run(value)
        rv.font.size = Pt(12)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # ÍNDICE (manual)
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "ÍNDICE", 1)
    toc_entries = [
        ("Capítulo 1 — Introducción", ""),
        ("  1.1 Contexto y motivación", ""),
        ("  1.2 Descripción del problema", ""),
        ("  1.3 Propuesta de solución", ""),
        ("  1.4 Objetivos del trabajo", ""),
        ("  1.5 Estructura del documento", ""),
        ("Capítulo 2 — Estado del Arte", ""),
        ("  2.1 Modelos de lenguaje de gran escala (LLMs)", ""),
        ("  2.2 Retrieval-Augmented Generation (RAG)", ""),
        ("  2.3 Arquitecturas de agentes con LLMs", ""),
        ("  2.4 Sistemas conversacionales en educación", ""),
        ("  2.5 Herramientas y tecnologías relevantes", ""),
        ("Capítulo 3 — Diseño del Sistema", ""),
        ("  3.1 Visión arquitectónica general", ""),
        ("  3.2 Diseño del Estado del grafo (LangGraph State)", ""),
        ("  3.3 Catálogo de intenciones (Intents)", ""),
        ("  3.4 Arquitectura de agentes", ""),
        ("  3.5 Diseño del flujo condicional del grafo", ""),
        ("  3.6 Diseño de la base de conocimiento", ""),
        ("  3.7 Diseño de la base de datos de alumnos", ""),
        ("  3.8 Diseño de los Guardrails", ""),
        ("Capítulo 4 — Desarrollo e Implementación (Fase 1)", ""),
        ("  4.1 Configuración del entorno de desarrollo", ""),
        ("  4.2 Construcción del corpus documental sintético", ""),
        ("  4.3 Diseño e implementación de la base de datos de alumnos", ""),
        ("  4.4 Implementación de la capa de acceso a datos", ""),
        ("  4.5 Configuración del proyecto", ""),
        ("  4.6 Verificación de la Fase 1", ""),
        ("Capítulo 5 — Planificación y Metodología", ""),
        ("  5.1 Metodología de desarrollo adoptada", ""),
        ("  5.2 Planificación del proyecto", ""),
        ("  5.3 Estado actual del proyecto", ""),
        ("Capítulo 6 — Conclusiones Parciales", ""),
        ("Referencias Bibliográficas", ""),
        ("Anexo A — Estructura del Proyecto", ""),
        ("Anexo B — Fragmentos de Código Relevantes", ""),
        ("Anexo C — Corpus Documental", ""),
    ]
    for entry, pg in toc_entries:
        p = doc.add_paragraph(entry)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)
            if not entry.startswith("  "):
                run.font.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # CAPÍTULO 1 — INTRODUCCIÓN
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Capítulo 1 — Introducción", 1)

    add_heading(doc, "1.1 Contexto y motivación", 2)
    add_body(doc, "Las instituciones universitarias afrontan hoy una tensión estructural entre la creciente demanda de orientación personalizada por parte del alumnado y los recursos humanos disponibles para satisfacerla. Los servicios de orientación académica operan en horarios de oficina, con capacidad de atención limitada y, en muchos casos, sin acceso inmediato al expediente completo del alumno que solicita ayuda. Esta situación genera cuellos de botella especialmente críticos en períodos de alta demanda: inicio de matrícula, época de exámenes o cierre de solicitudes de beca. Un alumno que a las 22:00 del día anterior a un plazo crítico necesita saber si puede anular una asignatura sin consumir convocatoria no tiene a quién preguntar.")
    add_body(doc, "Desde mi perspectiva como estudiante de Ingeniería Informática, he vivido de primera mano esta limitación. La falta de disponibilidad de información precisa y personalizada en los momentos de mayor necesidad —y la consiguiente angustia que genera— me llevó a plantearse si la tecnología actual, en concreto los grandes modelos de lenguaje combinados con técnicas de recuperación de información, podría ofrecer una solución efectiva. La motivación de este Trabajo de Fin de Grado nace precisamente de esa experiencia personal y de la convicción de que los recientes avances en inteligencia artificial generativa hacen por primera vez técnicamente viable un asistente universitario verdaderamente útil.")
    add_body(doc, "El problema no es nuevo. Los chatbots universitarios llevan años desplegándose en webs institucionales, pero su utilidad ha sido limitada porque respondían únicamente a preguntas predefinidas con respuestas genéricas, sin acceso a la documentación normativa actualizada y, desde luego, sin conocimiento alguno del expediente particular de cada alumno. La llegada de los modelos de lenguaje de gran escala (LLMs), y especialmente el paradigma Retrieval-Augmented Generation (RAG), abre la posibilidad de construir un sistema que combine lo mejor de ambos mundos: la fluidez conversacional de un LLM con la precisión factual de un sistema de recuperación de información basado en documentación institucional real.")
    add_body(doc, "He decidido abordar este proyecto con una arquitectura de múltiples agentes orquestados mediante LangGraph, una tecnología que permite modelar el flujo de procesamiento como un grafo de estados con transiciones condicionales. Esta elección refleja mi convicción de que la orientación universitaria es un proceso inherentemente complejo: no es simplemente buscar un dato en un documento, sino combinar la comprensión de la consulta, la recuperación de información relevante, el conocimiento del historial académico del alumno y la generación de una respuesta empática y personalizada. Cada uno de estos pasos merece un agente especializado, y LangGraph permite orquestarlos con la precisión y la trazabilidad que un sistema de uso real requiere.")

    add_heading(doc, "1.2 Descripción del problema", 2)
    add_body(doc, "El problema técnico central que aborda este TFG es el siguiente: cómo construir un sistema de asistencia conversacional que pueda responder consultas universitarias de forma personalizada, veraz y disponible en todo momento. La respuesta correcta a la pregunta «¿puedo anular la matrícula de Bases de Datos sin consumir convocatoria?» depende de varios factores simultáneos: la normativa de anulación de matrícula (que establece el plazo de la sexta semana lectiva), la asignatura concreta de la que se trate, y la situación académica del alumno (si ya ha consumido convocatorias previas). Ningún sistema puramente documental, ni ningún LLM sin acceso a datos externos, puede responder correctamente a esta pregunta en el caso general.")
    add_body(doc, "El problema se descompone en tres retos técnicos diferenciados. El primero es la recuperación precisa de información: dado que la base de conocimiento institucional puede contener decenas de documentos (reglamentos, planes de estudio, guías docentes, FAQs), el sistema debe ser capaz de identificar y recuperar los fragmentos más relevantes para cada consulta, sin contaminar la respuesta con información irrelevante. El segundo reto es la personalización: la respuesta debe tener en cuenta el expediente concreto del alumno —sus asignaturas aprobadas y pendientes, su nota media, su situación de beca— para que sea genuinamente útil y no meramente informativa en términos generales. El tercer reto es la seguridad: un sistema orientado a estudiantes, muchos de ellos en situaciones de estrés, debe ser capaz de detectar señales de crisis emocional y derivar apropiadamente, y debe negarse a responder consultas completamente ajenas al ámbito universitario.")
    add_body(doc, "A estos tres retos técnicos se añade un cuarto: la trazabilidad. En un contexto académico, el sistema debe ser capaz de justificar cada respuesta con referencias a los documentos fuente utilizados. El alumno debe poder verificar que la información que ha recibido proviene de la normativa oficial de la universidad, no de una alucinación del modelo de lenguaje. Esta exigencia descarta los enfoques basados únicamente en LLMs sin mecanismos de recuperación, y justifica el paradigma RAG como fundamento del sistema.")

    add_heading(doc, "1.3 Propuesta de solución", 2)
    add_body(doc, "He propuesto y diseñado un sistema de asistencia universitaria inteligente basado en el paradigma Retrieval-Augmented Generation (RAG) con orquestación de agentes mediante LangGraph. El sistema es accesible a través de una interfaz web construida con Streamlit, que actúa como capa de presentación, y tiene como núcleo un grafo de agentes LangGraph que orquesta el procesamiento de cada consulta de forma estructurada y trazable.")
    add_body(doc, "La Retrieval-Augmented Generation es un paradigma de procesamiento del lenguaje natural que combina la capacidad generativa de los LLMs con la recuperación de información de fuentes externas verificadas. En lugar de pedir al modelo que «recuerde» la normativa de anulación de matrícula —información que puede no estar en su conjunto de entrenamiento o estar desactualizada—, el sistema recupera dinámicamente los fragmentos relevantes de los documentos institucionales y los inyecta en el contexto del modelo antes de generar la respuesta. Esto elimina las alucinaciones factuales sobre información institucional y permite que la base de conocimiento se actualice sin necesidad de reentrenar el modelo.")
    add_body(doc, "LangGraph aporta la capacidad de modelar el flujo de procesamiento como un grafo de estados con transiciones condicionales. Cada nodo del grafo es un agente especializado con una responsabilidad concreta: el Router clasifica la intención del mensaje; el Guardrail evalúa su seguridad; el Retriever busca información en Azure AI Search; el agente Student Data consulta el expediente en SQLite; y el Generador sintetiza todo y produce la respuesta final. El estado del grafo —un TypedDict tipado que fluye por todos los nodos— garantiza que cada agente dispone exactamente de la información que necesita y que la respuesta final refleja todos los factores relevantes.")
    add_body(doc, "La combinación de RAG con orquestación multi-agente es especialmente adecuada para el problema universitario porque permite separar claramente las responsabilidades: la recuperación de documentación institucional (competencia del Retriever sobre Azure AI Search) y la recuperación de datos personales del alumno (competencia del agente Student Data sobre SQLite local) son flujos completamente distintos que merecen tratamiento diferenciado. El resultado es un asistente que puede decirle a un alumno concreto: «Dado que tienes aprobadas Matemáticas I y II con una media de 7.3, y estás en el tercer curso de Informática, te recomiendo matricularte en INF301 (IA Introducción), para la que ya cumples el prerrequisito de Probabilidad. Según el plan de estudios [Fuente: Plan Grado Informática, sección 3er curso], esta asignatura abre el camino a las optativas de Deep Learning y PLN de cuarto curso.»")

    add_heading(doc, "1.4 Objetivos del trabajo", 2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run("El objetivo principal de este TFG es diseñar, implementar y evaluar un sistema de asistencia universitaria inteligente basado en RAG y orquestación de agentes que sea capaz de responder consultas universitarias de forma personalizada, veraz y disponible 24/7.").font.size = Pt(12)
    doc.add_paragraph()
    obj_list = [
        "Diseñar una arquitectura de agentes multi-nodo con LangGraph que separe claramente las responsabilidades de clasificación, recuperación, personalización y generación.",
        "Implementar un sistema RAG sobre Azure AI Search con búsqueda híbrida (vectorial + semántica) capaz de recuperar fragmentos precisos de documentación universitaria.",
        "Construir un corpus documental sintético universitario suficientemente rico y estructurado para cubrir los casos de uso habituales de orientación académica.",
        "Diseñar e implementar una base de datos SQLite con el expediente académico de 50 alumnos sintéticos, con una capa de acceso basada en funciones predefinidas que impida la inyección SQL.",
        "Implementar un sistema de guardrails que detecte consultas fuera de alcance, contenido inapropiado y situaciones de crisis emocional, con respuestas predefinidas adecuadas para cada caso.",
        "Evaluar el sistema con el framework RAGAS, midiendo fidelidad, relevancia de la respuesta, precisión y recall del contexto recuperado, con objetivos cuantitativos definidos.",
        "Integrar el sistema completo en una interfaz web Streamlit con soporte de sesiones persistentes mediante el Checkpointer de LangGraph.",
        "Documentar de forma completa y reproducible el proceso de diseño, implementación y evaluación, de modo que el sistema pueda ser desplegado, mantenido y extendido por terceros.",
    ]
    for i, obj in enumerate(obj_list, 1):
        p = doc.add_paragraph(f"{i}. {obj}", style='List Number')
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.size = Pt(12)

    add_heading(doc, "1.5 Estructura del documento", 2)
    add_body(doc, "El presente documento se organiza en seis capítulos principales seguidos de tres anexos. El Capítulo 2 presenta el estado del arte en LLMs, RAG, arquitecturas de agentes y sistemas conversacionales educativos, con una tabla comparativa del stack tecnológico seleccionado. El Capítulo 3 describe el diseño completo del sistema: arquitectura general, estado del grafo, catálogo de intenciones, diseño de los cinco agentes, flujo condicional, diseño de la base de conocimiento y de la base de datos, y el sistema de guardrails. El Capítulo 4 documenta la implementación real de la Fase 1 del proyecto: corpus documental, base de datos SQLite, capa de acceso a datos y configuración del entorno, con referencias directas al código fuente generado. El Capítulo 5 describe la metodología de desarrollo adoptada y el plan de trabajo por fases. El Capítulo 6 presenta las conclusiones parciales al momento de esta entrega. Los anexos recogen la estructura de directorios del proyecto, los fragmentos de código más relevantes y el catálogo completo del corpus documental.")

    # ══════════════════════════════════════════════════════════════════════════
    # CAPÍTULO 2 — ESTADO DEL ARTE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Capítulo 2 — Estado del Arte", 1)

    add_heading(doc, "2.1 Modelos de lenguaje de gran escala (LLMs)", 2)
    add_body(doc, "La arquitectura Transformer, introducida por Vaswani et al. en 2017 con el artículo «Attention Is All You Need», supuso una revolución en el procesamiento del lenguaje natural. El mecanismo de atención multi-cabeza permite al modelo capturar dependencias de largo alcance entre tokens de forma paralela, superando las limitaciones de las arquitecturas recurrentes previas (LSTM, GRU). A partir de esta base, la escala de los modelos creció exponencialmente: GPT-2 (2019, 1.5B parámetros), GPT-3 (2020, 175B parámetros) y GPT-4 (2023, parámetros no revelados por OpenAI) demostraron que el incremento de parámetros y datos de entrenamiento produce saltos cualitativos en las capacidades del modelo, un fenómeno conocido como «emergent abilities».")
    add_body(doc, "Sin embargo, los LLMs puros presentan limitaciones estructurales relevantes para el caso de uso universitario. En primer lugar, el problema de las alucinaciones: el modelo genera texto estadísticamente plausible, pero no necesariamente factualmente correcto. Cuando un LLM afirma que «el plazo de anulación de matrícula es de ocho semanas», puede estar mezclando información de distintas universidades presentes en su corpus de entrenamiento. En segundo lugar, el conocimiento desactualizado: el modelo solo conoce la información presente en sus datos de entrenamiento hasta la fecha de corte, lo que lo hace inadecuado para responder sobre normativas actualizadas o calendarios académicos del curso en vigor. En tercer lugar, la falta de personalización: ningún LLM tiene acceso al expediente académico concreto del alumno que pregunta.")
    add_body(doc, "GPT-4o, el modelo que he seleccionado para este proyecto como motor generativo, supera a sus predecesores en comprensión contextual, capacidad de razonamiento y seguimiento de instrucciones complejas. Su versión desplegada mediante Azure OpenAI Service proporciona además las garantías de privacidad, SLA y residencia de datos que requiere un sistema universitario real. No obstante, estas ventajas no eliminan las limitaciones estructurales mencionadas, lo que refuerza la necesidad del enfoque RAG.")

    add_heading(doc, "2.2 Retrieval-Augmented Generation (RAG)", 2)
    add_body(doc, "El paradigma Retrieval-Augmented Generation fue formalizado por Lewis et al. en 2020 como solución a las limitaciones de los LLMs para tareas de knowledge-intensive NLP. La idea central es sencilla pero poderosa: en lugar de confiar en el conocimiento parametrizado del modelo, se conecta el LLM a una base de conocimiento externa que puede actualizarse de forma independiente. En el momento de generar una respuesta, el sistema recupera dinámicamente los fragmentos más relevantes de esa base y los incluye en el prompt, de modo que el modelo puede fundamentar su respuesta en información verificada.")
    add_body(doc, "El flujo RAG se descompone en dos fases. Durante la fase de indexación, los documentos fuente se dividen en fragmentos (chunks) de tamaño controlado, se genera un embedding vectorial de cada fragmento y se almacena en un índice vectorial. Durante la fase de inferencia, la consulta del usuario se transforma también en un embedding y se utiliza para recuperar los fragmentos más similares mediante búsqueda de vecinos aproximados. Los fragmentos recuperados se concatenan al prompt del LLM junto con la pregunta original, y el modelo genera una respuesta fundamentada en ese contexto enriquecido.")
    add_body(doc, "La búsqueda híbrida —combinación de búsqueda vectorial semántica y búsqueda léxica tradicional (BM25)— ha demostrado sistemáticamente superar a cualquiera de las dos por separado en benchmarks de recuperación de información. Azure AI Search implementa esta búsqueda híbrida con un perfil semántico configurable que reordena los resultados por relevancia conceptual, no solo por similitud de embeddings. He optado por esta aproximación en lugar de soluciones puramente vectoriales como Chroma o FAISS porque la búsqueda híbrida es más robusta ante consultas con terminología específica (nombres de asignaturas, códigos de artículos de reglamento) que pueden no estar bien representadas en el espacio de embeddings.")
    add_body(doc, "La evaluación de sistemas RAG es un problema activo de investigación. El framework RAGAS (Shahul Es et al., 2023) propone un conjunto de métricas sin necesidad de anotación humana: faithfulness (¿las afirmaciones de la respuesta están respaldadas por el contexto recuperado?), answer relevancy (¿la respuesta responde a la pregunta?), context precision (¿los chunks recuperados son relevantes?) y context recall (¿el contexto recuperado contiene la información necesaria?). He adoptado RAGAS como marco de evaluación principal, con umbrales objetivo de 0.85, 0.80, 0.75 y 0.70 respectivamente para cada métrica.")

    add_heading(doc, "2.3 Arquitecturas de agentes con LLMs", 2)
    add_body(doc, "LangChain, lanzado en 2022, introdujo el concepto de «cadena» (chain) como composición secuencial de llamadas a LLMs y herramientas. Las cadenas permitieron construir flujos de procesamiento más complejos que una sola llamada al modelo, pero adolecían de una limitación importante: el flujo era estático y no podía bifurcarse en función del contenido del mensaje. La llegada de LangGraph en 2024 resolvió esta limitación al modelar el flujo de procesamiento como un grafo de estados con transiciones condicionales, inspirado en la abstracción de máquinas de estado de Pregel (el motor de procesamiento de grafos de Google).")
    add_body(doc, "El concepto de agente en LangGraph es un nodo del grafo que recibe el estado actual, realiza alguna operación (llamada a un LLM, consulta a una herramienta externa, transformación de datos) y devuelve una actualización del estado. Los edges condicionales del grafo determinan a qué nodo se dirigirá el flujo a continuación en función del contenido del estado, permitiendo bifurcaciones, bucles y flujos completamente dinámicos. Esta abstracción es mucho más adecuada para el problema universitario que una cadena estática: la necesidad de consultar el expediente del alumno solo existe en ciertos tipos de consulta, y debe ser una decisión dinámica del grafo, no una operación siempre ejecutada.")
    add_body(doc, "He comparado LangGraph con dos alternativas principales antes de tomar esta decisión. AutoGen (Microsoft Research) adopta un enfoque de conversación entre agentes, donde múltiples LLMs se comunican entre sí para resolver tareas complejas. Aunque es poderoso para tareas de razonamiento profundo, introduce overhead de latencia y coste significativo para casos de uso conversacionales simples como preguntas frecuentes sobre matrícula. CrewAI sigue un paradigma similar pero con una abstracción más orientada a equipos de agentes con roles fijos. He elegido LangGraph porque su modelo de grafo de estados es la representación más natural para el flujo de procesamiento universitario: con estados bien definidos, condiciones de transición explícitas y trazabilidad completa del recorrido de cada mensaje a través del grafo.")

    add_heading(doc, "2.4 Sistemas conversacionales en educación", 2)
    add_body(doc, "Los chatbots universitarios no son una novedad. Sistemas como el de la Universidad de Georgia (POUNCE, 2019), el de Georgia Tech (Jill Watson, basado en IBM Watson) o los numerosos implementados sobre plataformas como Satisfi Labs o Engageware han demostrado la viabilidad del concepto y su impacto positivo en métricas de satisfacción estudiantil y reducción de carga del personal de administración. Sin embargo, todos ellos comparten una limitación estructural: están basados en sistemas de preguntas frecuentes predefinidas con respuestas estáticas, o en modelos de clasificación de intenciones seguidos de respuestas templates. Ninguno de ellos tiene acceso al expediente académico individual del alumno, ni puede razonar sobre normativas complejas con múltiples condiciones.")
    add_body(doc, "Los trabajos académicos más recientes exploran la aplicación de LLMs a entornos educativos, pero predominantemente para tareas de tutoría de contenido (ayudar al alumno a aprender) más que de orientación administrativa. Kasneci et al. (2023) ofrecen una revisión exhaustiva de las aplicaciones de ChatGPT y modelos similares en educación, señalando tanto su potencial como los riesgos de alucinaciones en contextos donde la precisión factual es crítica. Este TFG aborda específicamente ese riesgo mediante RAG sobre documentación institucional verificada.")
    add_body(doc, "Mi propuesta supera el estado del arte en el contexto universitario en tres dimensiones: la personalización mediante acceso al expediente real del alumno, la veracidad mediante RAG sobre documentación institucional actualizable sin reentrenamiento, y la seguridad mediante guardrails especializados para el contexto universitario (incluida la detección de crisis emocional con derivación al servicio de psicología). Ninguno de los sistemas publicados en la literatura combina estas tres dimensiones de forma integrada.")

    add_heading(doc, "2.5 Herramientas y tecnologías relevantes", 2)
    add_body(doc, "La siguiente tabla resume el proceso de selección tecnológica que he llevado a cabo para las componentes principales del sistema, con los criterios que he considerado determinantes:")
    add_table(doc,
        ["Componente", "Opción elegida", "Alternativas consideradas", "Razón principal de elección"],
        [
            ["Motor de búsqueda vectorial", "Azure AI Search", "Pinecone, Chroma, Weaviate, FAISS", "Búsqueda híbrida nativa, SLA empresarial, integración con Azure OpenAI"],
            ["Modelo generativo", "GPT-4o (Azure OpenAI)", "GPT-3.5-turbo, Llama 3, Mistral", "Mejor comprensión contextual y seguimiento de instrucciones complejas"],
            ["Embeddings", "text-embedding-3-large (3072d)", "text-embedding-ada-002, BGE-M3", "Mayor precisión semántica; benchmarks MTEB superiores"],
            ["Orquestación de agentes", "LangGraph", "AutoGen, CrewAI, LangChain LCEL", "Grafo de estados explícito, trazabilidad completa, condiciones dinámicas"],
            ["Base de datos expedientes", "SQLite", "PostgreSQL, MongoDB, DynamoDB", "Sin servidor, portabilidad total, suficiente para 50-1000 alumnos"],
            ["Interfaz web", "Streamlit", "Gradio, FastAPI+React, Dash", "Desarrollo rápido en Python puro, soporte nativo de chat y sesiones"],
            ["Evaluación RAG", "RAGAS", "TruLens, UpTrain, evaluación manual", "Métricas sin anotación, integración con LangChain, estándar de facto emergente"],
            ["Chunking", "tiktoken (cl100k_base)", "spaCy sentence splitter, RecursiveCharacterTextSplitter", "Conteo exacto de tokens del modelo, solapamiento preciso"],
        ]
    )

    # ══════════════════════════════════════════════════════════════════════════
    # CAPÍTULO 3 — DISEÑO DEL SISTEMA
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Capítulo 3 — Diseño del Sistema", 1)

    add_heading(doc, "3.1 Visión arquitectónica general", 2)
    add_body(doc, "He diseñado el sistema en tres capas bien diferenciadas que se comunican de forma asíncrona. La capa de presentación, implementada con Streamlit, gestiona la sesión del usuario, el historial de mensajes, la selección del alumno activo y la visualización de fuentes citadas. La capa de orquestación, implementada con LangGraph, contiene el grafo de agentes que procesa cada mensaje del usuario de forma estructurada. La capa de datos comprende Azure AI Search para el corpus documental vectorizado y SQLite para los expedientes académicos.")
    add_body(doc, "El flujo completo de un mensaje es el siguiente: (1) el usuario escribe en la interfaz Streamlit; (2) el mensaje se invoca sobre el grafo LangGraph con la configuración de sesión (thread_id); (3) el nodo Router analiza el mensaje y actualiza el State con el intent y los puntos clave; (4) el nodo Guardrail evalúa si la consulta es segura; (5) si el guardrail no se activa, el nodo Retriever ejecuta búsquedas en Azure AI Search; (6) si la consulta requiere datos del alumno y hay un alumno identificado, el nodo Student Data consulta SQLite; (7) el nodo Generador sintetiza toda la información del State y produce la respuesta final con fuentes citadas; (8) la respuesta se devuelve a Streamlit para su visualización.")
    add_code_block(doc,
"[Streamlit UI]\n"
"    │\n"
"    ▼  user_message + user_id + session_id\n"
"[LangGraph Graph]\n"
"    │\n"
"    ├─► [Router] → intent, key_points, requires_student_data\n"
"    │\n"
"    ├─► [Guardrail] → guardrail_triggered / safe_response\n"
"    │\n"
"    ├─► [Retriever] → retrieved_docs (Azure AI Search)\n"
"    │\n"
"    ├─► [Student Data] → student_record (SQLite) [condicional]\n"
"    │\n"
"    └─► [Generador] → final_response + sources + confidence\n"
"    │\n"
"    ▼\n"
"[Streamlit UI] ← respuesta + fuentes citadas")

    add_heading(doc, "3.2 Diseño del Estado del grafo (LangGraph State)", 2)
    add_body(doc, "He decidido modelar el estado del grafo como un TypedDict estrictamente tipado, siguiendo el patrón recomendado por LangGraph para garantizar la seguridad de tipos en tiempo de análisis estático. El estado es el único objeto que fluye por todos los nodos del grafo, actuando como contenedor centralizado de toda la información del ciclo de procesamiento. Cada nodo lee los campos que necesita y escribe únicamente los campos que le corresponden actualizar, sin acceso a información que no sea de su responsabilidad.")
    add_code_block(doc,
"class UniversityAssistantState(TypedDict):\n"
"    # Entrada del usuario\n"
"    user_message: str           # Mensaje original\n"
"    user_id: Optional[str]      # ID del alumno identificado\n"
"    session_id: str             # ID único de sesión\n"
"\n"
"    # Procesamiento\n"
"    intent: Optional[str]           # Intent clasificado por el Router\n"
"    key_points: List[str]           # Puntos clave extraídos\n"
"    requires_student_data: bool     # ¿Necesita expediente?\n"
"\n"
"    # Contexto recuperado\n"
"    retrieved_docs: List[Document]  # Chunks de Azure AI Search\n"
"    student_record: Optional[dict]  # Expediente del alumno\n"
"    search_queries: List[str]       # Queries ejecutadas\n"
"\n"
"    # Seguridad\n"
"    guardrail_triggered: bool\n"
"    guardrail_reason: Optional[str]\n"
"\n"
"    # Respuesta\n"
"    final_response: Optional[str]\n"
"    sources: List[str]\n"
"    confidence: float\n"
"\n"
"    # Historial\n"
"    message_history: List[dict]")

    add_heading(doc, "3.3 Catálogo de intenciones (Intents)", 2)
    add_body(doc, "He diseñado un catálogo de nueve intenciones mutuamente excluyentes que cubren el espacio de consultas universitarias habituales. La clasificación en exactamente un intent por mensaje es una decisión de diseño deliberada: simplifica la lógica de routing y permite que cada intent tenga estrategias de recuperación y generación optimizadas. La posible ambigüedad entre intents se resuelve en el prompt del Router mediante la instrucción de razonamiento paso a paso (Chain of Thought).")
    add_table(doc,
        ["Intent", "Descripción", "¿Requiere expediente?"],
        [
            ["ACADEMIC_ORIENTATION", "Itinerario, optativas recomendadas, planificación académica", "Sí"],
            ["ADMINISTRATIVE", "Plazos, matrícula, pagos, documentación", "A veces"],
            ["COURSE_INFO", "Información sobre asignaturas, guías docentes, profesores", "No"],
            ["REGULATIONS", "Normativas académicas, reglamento, permanencia", "No"],
            ["PROSPECTIVE_STUDENT", "Información para futuros alumnos: acceso, oferta, precios", "No"],
            ["SCHOLARSHIPS", "Becas disponibles, requisitos, plazos de solicitud", "A veces"],
            ["EMOTIONAL_SUPPORT", "Estrés, problemas personales → derivar a psicología", "No"],
            ["OUT_OF_SCOPE", "Consulta completamente ajena al ámbito universitario", "No"],
            ["GREETING", "Saludo inicial o mensaje de cortesía", "No"],
        ]
    )

    add_heading(doc, "3.4 Arquitectura de agentes", 2)

    add_heading(doc, "Agente 1 — Router (clasificador de intenciones)", 2)
    add_body(doc, "El Router es el primer nodo que procesa cualquier mensaje entrante. Su responsabilidad exclusiva es analizar el mensaje del usuario y el historial reciente de la conversación para producir tres outputs: el intent (clasificado en el catálogo de nueve valores), los key_points (entidades y temas clave extraídos del mensaje) y el flag requires_student_data (si la respuesta óptima requiere acceso al expediente del alumno). El Router utiliza GPT-4o con el ROUTER_SYSTEM_PROMPT definido, que instruye al modelo a responder únicamente con un JSON válido de estructura fija. No genera ningún texto para el usuario.")
    add_body(doc, "He diseñado el Router para que opere con razonamiento explícito paso a paso (Chain of Thought): el prompt lo instruye a identificar primero el tema principal, luego las entidades clave, luego determinar si se necesita el expediente, y finalmente clasificar el intent. Este enfoque reduce los errores de clasificación en consultas ambiguas. Analicé alternativas como clasificadores fine-tuned o embeddings de similitud de coseno sobre ejemplos de cada intent, pero los descartéporque el LLM con CoT ofrece mayor flexibilidad ante consultas que no encajan perfectamente en ningún ejemplo preexistente.")

    add_heading(doc, "Agente 2 — Guardrail (seguridad y límites)", 2)
    add_body(doc, "El Guardrail actúa inmediatamente después del Router, antes de que se ejecute ninguna búsqueda ni consulta a la base de datos. Evalúa el intent y el mensaje original para determinar si la consulta debe ser bloqueada o derivada antes de procesarse. He definido tres tipos de guardrail: out_of_scope (el intent es OUT_OF_SCOPE o la consulta es claramente ajena al ámbito universitario), emotional_crisis (detección de señales de crisis como «no puedo más», «me hago daño», «quiero dejarlo todo»), e inappropriate (contenido ofensivo o solicitudes de información sensible).")
    add_body(doc, "Si el guardrail se activa (guardrail_triggered = True), el grafo salta directamente al nodo Generador con una respuesta predefinida almacenada en el campo guardrail_reason. La respuesta para crisis emocional incluye obligatoriamente los datos de contacto del Servicio de Psicología (psicologia@universidad.es, 900 456 789). He tomado la decisión de no procesar la consulta con el Retriever cuando el guardrail se activa, tanto por razones de eficiencia como para evitar que una consulta manipuladora obtenga información del sistema aunque esté formulada como crisis emocional.")

    add_heading(doc, "Agente 3 — Retriever (búsqueda en Azure AI Search)", 2)
    add_body(doc, "El Retriever es el agente responsable de recuperar información relevante del corpus documental. A partir de los key_points producidos por el Router, construye una o varias queries de búsqueda optimizadas y las ejecuta contra el índice Azure AI Search configurado con búsqueda híbrida. El número máximo de documentos recuperados está configurado mediante la constante MAX_RETRIEVED_DOCS (por defecto 5), y el umbral de relevancia mínimo mediante RETRIEVAL_SCORE_THRESHOLD (por defecto 0.7).")
    add_body(doc, "He diseñado el Retriever para que ejecute múltiples queries cuando la consulta contiene varios temas distintos. Por ejemplo, una pregunta sobre «qué optativas de IA me recomiendas para mi situación» generaría dos queries: una sobre el plan de estudios de las optativas de IA disponibles y otra sobre los prerrequisitos de esas asignaturas. Los resultados se deduplican por chunk_id antes de pasarlos al Generador. El agente escribe en el State los campos retrieved_docs y search_queries (para trazabilidad).")

    add_heading(doc, "Agente 4 — Student Data (consulta de expediente)", 2)
    add_body(doc, "El agente Student Data se activa únicamente cuando se cumplen dos condiciones simultáneas: requires_student_data es True (determinado por el Router) y hay un user_id válido en el State (el alumno está identificado en la interfaz). Su implementación es deliberadamente simple: llama a la función get_full_student_record(user_id) del módulo sqlite_query.py y almacena el resultado en el campo student_record del State.")
    add_body(doc, "He tomado la decisión crítica de no implementar Text-to-SQL en este agente. Un enfoque Text-to-SQL permitiría al LLM generar consultas SQL dinámicamente, lo que conllevaría riesgos de inyección SQL y resultados impredecibles. En su lugar, he implementado un conjunto fijo de ocho funciones Python con queries parametrizadas y seguras, que cubren todos los casos de uso identificados para el sistema. Este enfoque sacrifica cierta flexibilidad en favor de seguridad y predecibilidad, una compensación que considero completamente justificada en un sistema que maneja datos académicos de alumnos reales.")

    add_heading(doc, "Agente 5 — Generador de respuesta", 2)
    add_body(doc, "El Generador es el último nodo del grafo y el único que produce texto para el usuario. Recibe el State completo (intent, retrieved_docs, student_record, guardrail_triggered, historial) e invoca GPT-4o con el GENERATOR_SYSTEM_PROMPT, que define la persona del asistente (psicopedagogo universitario empático y preciso), las reglas de citación de fuentes y las instrucciones de personalización cuando hay expediente del alumno disponible.")
    add_body(doc, "He diseñado el Generador para que la longitud de la respuesta sea proporcional a la complejidad de la consulta: un saludo merece un párrafo; una orientación de itinerario merece tres a cinco párrafos con referencias a asignaturas concretas y fuentes citadas. La instrucción explícita de nunca afirmar nada que no esté respaldado por los documentos recuperados, y de decirlo explícitamente si la información no está disponible, es la principal salvaguarda contra las alucinaciones factuales. El Generador escribe final_response, sources y confidence en el State.")

    add_heading(doc, "3.5 Diseño del flujo condicional del grafo", 2)
    add_body(doc, "El grafo LangGraph implementa las siguientes condiciones de transición entre nodos:")
    add_code_block(doc,
"START → Router  (siempre)\n"
"\n"
"Router → Guardrail  (siempre)\n"
"\n"
"Guardrail → Generador   SI guardrail_triggered == True\n"
"Guardrail → Retriever   SI guardrail_triggered == False\n"
"\n"
"Retriever → Student Data  SI requires_student_data == True\n"
"                           Y user_id is not None\n"
"Retriever → Generador     EN CASO CONTRARIO\n"
"\n"
"Student Data → Generador  (siempre)\n"
"\n"
"Generador → END  (siempre)")
    add_body(doc, "Este diseño garantiza que el Retriever siempre se ejecuta para consultas seguras (la recuperación de contexto documental es siempre beneficiosa), mientras que la consulta de expediente es estrictamente condicional. La separación entre el nodo Guardrail y el nodo Generador permite que el Generador utilice las respuestas predefinidas del guardrail directamente desde el State, sin necesidad de lógica adicional.")

    add_heading(doc, "3.6 Diseño de la base de conocimiento", 2)
    add_body(doc, "He diseñado la base de conocimiento en dos capas con responsabilidades complementarias. Azure AI Search almacena el corpus documental institucional: normativas, planes de estudio, guías docentes, FAQs y procedimientos. Es la fuente de verdad para preguntas generales sobre la universidad, accesible a cualquier usuario independientemente de si está identificado. SQLite almacena los expedientes académicos individuales: notas, asignaturas, becas. Es la fuente de personalización, accesible únicamente para el alumno identificado con su propio expediente.")
    add_body(doc, "Para el chunking del corpus documental, he decidido utilizar fragmentos de 512 tokens medidos con el tokenizador cl100k_base de tiktoken, con un solapamiento de 50 tokens entre chunks consecutivos. La elección de 512 tokens equilibra la densidad de información (suficiente contexto para que el chunk sea autosuficiente) con la precisión de la recuperación (chunks demasiado largos difuminan la señal de relevancia). El solapamiento de 50 tokens evita cortes en medio de una idea cuando el artículo de reglamento o la sección del plan de estudios supera los 512 tokens. Los separadores de chunk priorizan los saltos de párrafo sobre los puntos de fin de frase, y nunca cortan en medio de una oración.")
    add_body(doc, "Cada chunk indexado incluye metadatos estructurados que permiten citar la fuente con precisión: id (identificador único del chunk), content (texto), title (título del documento), category (normativa/plan_estudio/guia_docente/faq/procedimiento), source_file (nombre del archivo), page_number, section (título de la sección o artículo), y degree_relevance (lista de grados a los que aplica el documento). Este diseño de metadatos permite al Generador producir citas del tipo «[Fuente: Reglamento Académico General, Art. 4 – Convocatorias, pág. 3]».")

    add_heading(doc, "3.7 Diseño de la base de datos de alumnos", 2)
    add_body(doc, "He diseñado el esquema SQLite con tres tablas normalizadas que representan las tres dimensiones del expediente académico: el perfil del alumno (students), el historial de asignaturas (subject_records) y las becas (scholarships). La separación en tres tablas, en lugar de una única tabla desnormalizada, respeta los principios de la tercera forma normal y facilita las consultas específicas de cada tipo de información.")
    add_code_block(doc,
"CREATE TABLE students (\n"
"    id                TEXT PRIMARY KEY,   -- 'ALU001'\n"
"    name              TEXT NOT NULL,\n"
"    email             TEXT UNIQUE,\n"
"    degree            TEXT,\n"
"    year              INTEGER,\n"
"    gpa               REAL,\n"
"    credits_completed INTEGER,\n"
"    credits_total     INTEGER,\n"
"    status            TEXT,              -- 'active', 'at_risk', 'excellent'\n"
"    enrolled_year     INTEGER\n"
");\n\n"
"CREATE TABLE subject_records (\n"
"    id           INTEGER PRIMARY KEY AUTOINCREMENT,\n"
"    student_id   TEXT REFERENCES students(id),\n"
"    subject_code TEXT,\n"
"    subject_name TEXT,\n"
"    credits      INTEGER,\n"
"    grade        REAL,\n"
"    semester     TEXT,\n"
"    status       TEXT  -- 'passed', 'failed', 'enrolled', 'pending'\n"
");\n\n"
"CREATE TABLE scholarships (\n"
"    id         INTEGER PRIMARY KEY AUTOINCREMENT,\n"
"    student_id TEXT REFERENCES students(id),\n"
"    type       TEXT,  -- 'MEC', 'universidad', 'movilidad'\n"
"    amount     REAL,\n"
"    year       INTEGER,\n"
"    status     TEXT   -- 'active', 'expired', 'pending'\n"
");")

    add_heading(doc, "3.8 Diseño de los Guardrails", 2)
    add_body(doc, "He incluido un sistema de guardrails como componente obligatorio porque el sistema está orientado a estudiantes universitarios, un colectivo especialmente vulnerable en términos de salud mental. El contexto académico —presión de exámenes, riesgo de baja académica, problemas económicos de beca— puede detonar situaciones de crisis emocional que requieren derivación inmediata a profesionales humanos. Un sistema de asistencia que responde únicamente con información sobre normativas en ese contexto sería éticamente inadecuado.")
    add_body(doc, "Los tres tipos de guardrail implementados son: out_of_scope (la consulta no tiene relación con el ámbito universitario; respuesta educada que redirige al propósito del asistente), emotional_crisis (detección de patrones lingüísticos de angustia severa; respuesta empática con datos de contacto del Servicio de Psicología: psicologia@universidad.es, 900 456 789, lunes-viernes 9h-18h) e inappropriate (contenido ofensivo o solicitudes de información que podría comprometer la privacidad de otros alumnos; respuesta firme pero respetuosa). El guardrail opera sobre el texto del mensaje original, no sobre el intent clasificado por el Router, para capturar casos donde el formulado de la consulta contiene señales de crisis aunque el intent parezca académico.")

    # ══════════════════════════════════════════════════════════════════════════
    # CAPÍTULO 4 — DESARROLLO E IMPLEMENTACIÓN
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Capítulo 4 — Desarrollo e Implementación (Fase 1)", 1)

    add_heading(doc, "4.1 Configuración del entorno de desarrollo", 2)
    add_body(doc, "He configurado el entorno de desarrollo sobre Windows 10 Pro con Python 3.10. El control de versiones se gestiona con Git, inicializado en la raíz del proyecto con un archivo .gitignore que excluye explícitamente el archivo .env (nunca debe subirse al repositorio), los archivos .db de base de datos, los directorios de logs, cachés de Python y entornos virtuales.")
    add_body(doc, "La estructura de directorios del proyecto, tal como existe actualmente en disco, es la siguiente:")
    add_code_block(doc,
"tfg-universidad-rag/\n"
"├── .env.example          # Template de variables de entorno (vacías)\n"
"├── .gitignore            # Exclusiones Git\n"
"├── CLAUDE.md             # Instrucciones para Claude Code\n"
"├── README.md             # Descripción y guía de instalación\n"
"├── requirements.txt      # Dependencias Python (13 paquetes)\n"
"├── app/\n"
"│   ├── __init__.py\n"
"│   └── components/\n"
"│       └── __init__.py   # (pendiente Fase 4)\n"
"├── data/\n"
"│   ├── corpus/\n"
"│   │   ├── normativas/   (3 archivos .md)\n"
"│   │   ├── planes_estudio/ (3 archivos .md)\n"
"│   │   ├── guias_docentes/ (8 archivos .md)\n"
"│   │   ├── faqs/         (4 archivos .md)\n"
"│   │   └── procedimientos/ (4 archivos .md)\n"
"│   └── database/\n"
"│       └── students.db   # 50 alumnos, 1114 registros de asignaturas\n"
"├── docs/\n"
"│   └── diagrams/         # (pendiente Fase 3)\n"
"├── logs/\n"
"│   └── fase1_20260423.txt\n"
"├── scripts/\n"
"│   ├── generate_students.py\n"
"│   └── verify_phase1.py\n"
"├── src/\n"
"│   ├── agents/           # (pendiente Fase 3)\n"
"│   ├── config/\n"
"│   │   ├── azure_config.py\n"
"│   │   └── settings.py\n"
"│   ├── graph/            # (pendiente Fase 3)\n"
"│   ├── prompts/          # (pendiente Fase 3)\n"
"│   └── tools/\n"
"│       └── sqlite_query.py\n"
"└── tests/                # (pendiente Fase 5)")

    add_heading(doc, "4.2 Construcción del corpus documental sintético", 2)
    add_body(doc, "He optado por construir un corpus documental completamente sintético en lugar de utilizar documentación real de ninguna universidad concreta. Esta decisión tiene dos justificaciones: una práctica (evitar problemas de derechos de autor y de datos sensibles de alumnos reales) y una metodológica (el corpus sintético me permite controlar exactamente qué información contiene, facilitando la evaluación posterior con RAGAS, donde el ground truth puede derivarse directamente de los propios documentos generados).")
    add_body(doc, "He generado 22 archivos Markdown estructurados en cinco categorías, utilizando herramientas de asistencia para la redacción del contenido pero tomando personalmente todas las decisiones de estructura, cobertura temática y criterios de calidad. Los documentos están redactados en el contexto de la «Universidad Demo», una institución ficticia cuyos datos (notas de corte, precios de matrícula, artículos de reglamento) son coherentes entre sí aunque no corresponden a ninguna universidad real. La siguiente tabla resume el corpus construido:")
    add_table(doc,
        ["Categoría", "Archivo", "Palabras aprox."],
        [
            ["normativa", "reglamento_academico.md", "1.544"],
            ["normativa", "normativa_permanencia.md", "1.053"],
            ["normativa", "politica_evaluacion.md", "1.134"],
            ["plan_estudio", "grado_informatica.md", "1.235"],
            ["plan_estudio", "grado_ade.md", "1.146"],
            ["plan_estudio", "grado_derecho.md", "1.178"],
            ["guia_docente", "guia_programacion_i.md", "617"],
            ["guia_docente", "guia_estructuras_datos.md", "606"],
            ["guia_docente", "guia_bases_datos.md", "624"],
            ["guia_docente", "guia_redes.md", "616"],
            ["guia_docente", "guia_ia_introduccion.md", "624"],
            ["guia_docente", "guia_matematicas_i.md", "626"],
            ["guia_docente", "guia_calculo.md", "587"],
            ["guia_docente", "guia_estadistica.md", "609"],
            ["faq", "faq_matriculacion.md", "1.055"],
            ["faq", "faq_becas.md", "1.115"],
            ["faq", "faq_admision.md", "967"],
            ["faq", "faq_examenes.md", "1.042"],
            ["procedimiento", "guia_orientacion_academica.md", "798"],
            ["procedimiento", "protocolo_apoyo_riesgo.md", "1.094"],
            ["procedimiento", "servicios_universidad.md", "1.270"],
            ["procedimiento", "calendario_academico.md", "856"],
            ["TOTAL", "22 archivos", "20.396"],
        ]
    )
    add_body(doc, "He aplicado criterios de calidad rigurosos a todos los documentos. Cada documento sigue un formato uniforme con cabecera de metadatos (título, universidad, fecha, categoría), secciones claramente tituladas con encabezados Markdown nivel 2 y 3, y contenido suficientemente denso para que el RAG pueda recuperar fragmentos con contexto autosuficiente. El Reglamento Académico, por ejemplo, incluye obligatoriamente artículos sobre el número máximo de convocatorias (4 ordinarias, posibilidad de 5ª extraordinaria), los requisitos de permanencia (50% de créditos en el primer año), el sistema de calificación (0-10, aprobado con 5.0), el plazo de anulación de matrícula (sexta semana lectiva) y el procedimiento de reclamación de calificaciones (5 días hábiles tras publicación). Estos son precisamente los temas sobre los que los alumnos formulan más consultas en los servicios de orientación.")

    add_heading(doc, "4.3 Diseño e implementación de la base de datos de alumnos", 2)
    add_body(doc, "He diseñado e implementado el script scripts/generate_students.py que crea y puebla la base de datos students.db con datos sintéticos pero coherentes. El script utiliza exclusivamente la biblioteca estándar sqlite3 de Python y el módulo random (con semilla fija para reproducibilidad). La generación es determinista: la misma semilla produce siempre los mismos 50 alumnos.")
    add_body(doc, "He distribuido los 50 alumnos en tres grados y cuatro cursos con la siguiente distribución: 20 alumnos de Ingeniería Informática, 15 de ADE y 15 de Derecho. Por estado académico: 10 alumnos excelentes (GPA > 8.5), 30 activos (GPA 5.0-8.5) y 10 en riesgo académico (GPA < 5.0). He definido cuatro alumnos especiales con perfiles concretos para facilitar las pruebas del sistema: ALU001 (Informática 4º, GPA 9.1, status excellent, matriculado en TFG), ALU015 (Informática 2º, GPA 4.2, status at_risk, asignaturas suspensas repetidas), ALU030 (ADE 3º, GPA 7.8, status active, beca MEC activa de 1.500 €) y ALU045 (Derecho 1º, GPA 5.5, status active, sin historial de aprobadas).")
    add_body(doc, "Las estadísticas reales de la base de datos generada son: tabla students con 50 filas, tabla subject_records con 1.114 filas (654 pasadas, 250 matriculadas en el semestre actual, 166 pendientes y 44 suspensas), y tabla scholarships con 24 filas (14 becas MEC activas y 10 becas de excelencia de la propia universidad).")

    add_heading(doc, "4.4 Implementación de la capa de acceso a datos (sqlite_query.py)", 2)
    add_body(doc, "He implementado el módulo src/tools/sqlite_query.py con ocho funciones Python predefinidas que constituyen la única interfaz entre el sistema y la base de datos de alumnos. La decisión de no implementar Text-to-SQL es una decisión de seguridad: si el LLM pudiera escribir SQL dinámico, un usuario malicioso podría formular mensajes diseñados para extraer información de otros alumnos o para corromper la base de datos. Las funciones predefinidas utilizan exclusivamente parámetros vinculados (parameterized queries con ?), eliminando completamente el riesgo de inyección SQL.")
    add_body(doc, "Las ocho funciones implementadas son: get_student_profile (perfil básico), get_subject_records (todas las asignaturas), get_pending_subjects (asignaturas pendientes o suspensas), get_passed_subjects (asignaturas superadas con nota), get_current_enrollments (asignaturas matriculadas en el semestre en curso), get_scholarships (becas activas o pendientes), get_academic_standing (créditos completados, porcentaje de avance, situación de riesgo calculada), y get_full_student_record (consolidación de todas las anteriores en un único dict estructurado, que es el único método que invoca el agente Student Data).")
    add_body(doc, "El siguiente fragmento muestra el patrón de implementación utilizado en todas las funciones, con context manager para gestión automática de conexiones, row_factory para acceso por nombre de campo y manejo de errores que nunca lanza excepciones al llamador:")
    add_code_block(doc,
"def get_student_profile(student_id: str) -> dict:\n"
"    logger.debug(\"get_student_profile(%s)\", student_id)\n"
"    try:\n"
"        with sqlite3.connect(_get_db_path()) as conn:\n"
"            conn.row_factory = sqlite3.Row\n"
"            row = conn.execute(\n"
"                \"\"\"SELECT id, name, email, degree, year, gpa,\n"
"                          credits_completed, credits_total, status\n"
"                   FROM students WHERE id = ?\"\"\",\n"
"                (student_id,),\n"
"            ).fetchone()\n"
"            if row is None:\n"
"                return {\"error\": f\"Alumno '{student_id}' no encontrado.\"}\n"
"            return dict(row)\n"
"    except sqlite3.Error as exc:\n"
"        logger.error(\"Error en get_student_profile: %s\", exc)\n"
"        return {\"error\": f\"Error de base de datos: {exc}\"}")

    add_heading(doc, "4.5 Configuración del proyecto", 2)
    add_body(doc, "He organizado la configuración del sistema en tres niveles. El archivo .env.example (nunca el .env real) define las diecisiete variables de entorno requeridas con valores vacíos o comentarios explicativos: credenciales Azure OpenAI y Azure AI Search, ruta de la base de datos SQLite, parámetros del sistema (MAX_RETRIEVED_DOCS, RETRIEVAL_SCORE_THRESHOLD), modo debug y flag del Checkpointer. El archivo src/config/settings.py centraliza las constantes del sistema, leyéndolas del entorno con valores por defecto seguros que permiten importar el módulo sin credenciales Azure. El archivo src/config/azure_config.py define las funciones de inicialización de clientes Azure (get_search_client, get_index_client, get_openai_client), que siguen un patrón de inicialización lazy: el módulo se importa sin error, pero las funciones lanzan ValueError descriptivos si se invocan sin las credenciales configuradas.")
    add_body(doc, "La separación entre settings.py (constantes del sistema) y azure_config.py (inicialización de clientes) es una decisión de diseño que facilita las pruebas unitarias: los tests que no requieren Azure pueden importar settings sin necesidad de un .env completo, mientras que los tests de integración configuran el entorno antes de invocar los clientes Azure.")

    add_heading(doc, "4.6 Verificación de la Fase 1", 2)
    add_body(doc, "He implementado el script scripts/verify_phase1.py como herramienta de verificación automatizada que comprueba la completitud y corrección de todos los artefactos de la Fase 1. El script realiza 21 checks organizados en tres bloques correspondientes a los tres bloques de trabajo de la fase. A continuación se muestra el output real de la última ejecución:")
    add_code_block(doc,
"══════════════════════════════════════════\n"
"  VERIFICACIÓN FASE 1 — TFG Universidad  \n"
"══════════════════════════════════════════\n\n"
"[BLOQUE 1.1 — Configuración]\n"
"  ✓ .gitignore existe\n"
"  ✓ .env.example existe y tiene 12+ variables\n"
"  ✓ requirements.txt existe con 12+ dependencias\n"
"  ✓ README.md existe (>500 palabras)\n"
"  ✓ CLAUDE.md existe\n"
"  ✓ src/config/settings.py importable sin credenciales\n"
"  ✓ src/config/azure_config.py importable sin credenciales\n\n"
"[BLOQUE 1.2 — Corpus]\n"
"  ✓ Normativas: 3/3 archivos presentes\n"
"  ✓ Planes de estudio: 3/3 archivos presentes\n"
"  ✓ Guías docentes: 8/8 archivos presentes\n"
"  ✓ FAQs: 4/4 archivos presentes\n"
"  ✓ Procedimientos: 4/4 archivos presentes\n"
"  ✓ Total palabras corpus: 20396 (mínimo requerido: 15000)\n\n"
"[BLOQUE 1.3 — Base de datos]\n"
"  ✓ data/database/students.db existe\n"
"  ✓ Tabla students: 50 filas\n"
"  ✓ Tabla subject_records: 1114 filas\n"
"  ✓ Tabla scholarships: 24 filas\n"
"  ✓ ALU001 existe con status='excellent'\n"
"  ✓ ALU015 existe con status='at_risk'\n"
"  ✓ sqlite_query.py: 8 funciones importadas correctamente\n"
"  ✓ get_full_student_record('ALU001') devuelve dict con 6 claves\n\n"
"══════════════════════════════════════════\n"
"  RESULTADO: 21/21 checks PASS\n"
"══════════════════════════════════════════")

    # ══════════════════════════════════════════════════════════════════════════
    # CAPÍTULO 5 — PLANIFICACIÓN Y METODOLOGÍA
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Capítulo 5 — Planificación y Metodología", 1)

    add_heading(doc, "5.1 Metodología de desarrollo adoptada", 2)
    add_body(doc, "He adoptado un enfoque de desarrollo iterativo organizado en cinco fases secuenciales, cada una con criterios de aceptación verificables y un script de verificación automatizado antes de avanzar a la siguiente. Esta metodología es especialmente adecuada para proyectos de inteligencia artificial porque permite validar hipótesis de forma incremental: la Fase 1 establece los datos y la configuración sin necesidad de servicios cloud, la Fase 2 añade la capa de recuperación verificando la calidad de la indexación, la Fase 3 integra el sistema completo de agentes y la Fase 4 añade la interfaz visual. La Fase 5 realiza la evaluación formal con RAGAS.")
    add_body(doc, "He definido explícitamente el criterio de no avanzar a la siguiente fase sin que la actual esté al 100% verificada. Esta restricción, que puede parecer restrictiva, evita el error habitual en proyectos de IA de construir capas superiores sobre fundamentos inestables. La calidad del corpus documental y de la base de datos de alumnos condiciona directamente la calidad de las respuestas del sistema; un error en esa base se manifestaría de forma opaca en fases superiores, siendo difícil de diagnosticar. Las dependencias entre servicios externos (Azure) crean puntos de bloqueo planificados donde el trabajo de implementación puede continuar en paralelo con la configuración de los recursos cloud.")

    add_heading(doc, "5.2 Planificación del proyecto", 2)
    add_table(doc,
        ["Fase", "Objetivo principal", "Dependencias", "Estado"],
        [
            ["Fase 1 — Cimientos y datos", "Corpus, SQLite, configuración y scripts sin Azure", "Ninguna", "✅ Completada"],
            ["Fase 2 — Azure AI Search", "Índice vectorial, indexación del corpus, tool de búsqueda", "Credenciales Azure", "⏳ Pendiente"],
            ["Fase 3 — Agentes y grafo", "State, prompts, 5 agentes, grafo LangGraph completo", "Fase 2 operativa", "⏳ Pendiente"],
            ["Fase 4 — Interfaz Streamlit", "Componentes UI, integración grafo-Streamlit, sesiones", "Fase 3 operativa", "⏳ Pendiente"],
            ["Fase 5 — Tests y evaluación", "Dataset RAGAS, evaluación, tests unitarios, cierre", "Sistema completo", "⏳ Pendiente"],
        ]
    )

    add_heading(doc, "5.3 Estado actual del proyecto", 2)
    add_body(doc, "Al momento de esta entrega (Entrega 1, abril 2025), la Fase 1 está completada al 100% con todos los checks de verificación en verde. Los artefactos disponibles son: 22 documentos del corpus sintético (20.396 palabras), la base de datos SQLite con 50 alumnos y 1.114 registros de asignaturas, el módulo de acceso a datos con 8 funciones implementadas y probadas, los archivos de configuración del proyecto y el script de verificación automatizada.")
    add_body(doc, "El principal bloqueo actual para avanzar a la Fase 2 es la provisión de credenciales Azure (Azure OpenAI Service y Azure AI Search). Una vez disponibles, el trabajo de indexación del corpus (Fase 2) puede completarse en una sesión de trabajo de 2-3 horas, pues los scripts y el corpus ya están preparados. Los riesgos identificados son: (1) Coste de las APIs Azure: el uso de GPT-4o y text-embedding-3-large tiene un coste por token que debe monitorizarse para no exceder el presupuesto del proyecto. Mitigación: configurar alertas de coste en Azure Portal. (2) Calidad del corpus: el corpus sintético puede no cubrir todos los escenarios de consulta reales. Mitigación: el sistema de evaluación RAGAS identificará las lagunas de cobertura antes de la entrega final. (3) Latencia del sistema multi-agente: la cadena de cinco agentes puede introducir latencias perceptibles. Mitigación: medición de latencia en Fase 3 con optimización de prompts si es necesario.")

    # ══════════════════════════════════════════════════════════════════════════
    # CAPÍTULO 6 — CONCLUSIONES PARCIALES
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Capítulo 6 — Conclusiones Parciales", 1)

    add_body(doc, "Al concluir la Fase 1 del proyecto, puedo afirmar que se ha establecido un cimiento técnico sólido y verificable sobre el que construir el resto del sistema. He completado el diseño arquitectónico completo del sistema, la base de conocimiento universitaria sintética con más de 20.000 palabras de contenido estructurado, la base de datos de expedientes académicos con 50 alumnos distribuidos en tres grados y cuatro perfiles de rendimiento, y la capa de acceso a datos con seguridad de tipos y sin riesgo de inyección SQL. Todos estos componentes han superado la verificación automatizada con 21/21 checks en verde.")
    add_body(doc, "Los desafíos encontrados durante la Fase 1 han sido principalmente de naturaleza técnica y han contribuido a decisiones de diseño mejores que las originalmente previstas. El más significativo fue la gestión de la asignación de IDs en la generación de alumnos sintéticos: el script inicial colisionaba entre los alumnos «especiales» de prueba (ALU001, ALU015, ALU030, ALU045) y los alumnos generados por rango, produciendo un error de restricción de unicidad en SQLite. La resolución de este problema me llevó a rediseñar la lógica de asignación de IDs para separar explícitamente el pool de IDs libres del conjunto de IDs reservados para alumnos de prueba, produciendo un código más robusto y mantenible. El segundo desafío fue la codificación de caracteres Unicode en el terminal Windows con codificación cp1252, que impedía mostrar los símbolos ✓ y ✗ del script de verificación. Se resolvió añadiendo el flag -X utf8 de Python y configurando stdout explícitamente como UTF-8.")
    add_body(doc, "Reflexionando sobre las decisiones de diseño tomadas, mantendría la decisión de usar funciones predefinidas en lugar de Text-to-SQL para el acceso a la base de datos de alumnos: es la decisión correcta desde el punto de vista de la seguridad, aunque reduce la flexibilidad de consulta. También mantendría la decisión de construir un corpus sintético en lugar de reutilizar documentación real: el control sobre el contenido facilita enormemente la evaluación con RAGAS. Si pudiera revisitar alguna decisión, consideraría incluir más documentos de guías docentes desde el principio (actualmente solo hay 8; cubrir las 35-40 asignaturas obligatorias de los tres grados daría una cobertura RAG mucho más completa). Sin embargo, esta ampliación puede realizarse en cualquier momento antes de la indexación en la Fase 2, sin modificar ningún componente del sistema.")
    add_body(doc, "Las expectativas para las fases siguientes son positivas. La Fase 2 está bien acotada: una vez disponibles las credenciales Azure, el script de indexación puede implementarse y ejecutarse en pocas horas. La Fase 3 (agentes y grafo) es la más compleja y la que más incertidumbre técnica lleva, especialmente en el diseño de los prompts de cada agente para producir consistentemente el formato JSON esperado. He destinado el grueso del tiempo de trabajo previsto a esa fase. La evaluación RAGAS de la Fase 5 permitirá cuantificar objetivamente la calidad del sistema y justificar las decisiones de diseño con datos empíricos, lo que consideroclave para la calidad académica del trabajo.")

    # ══════════════════════════════════════════════════════════════════════════
    # REFERENCIAS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Referencias Bibliográficas", 1)

    refs = [
        "Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, Ł., & Polosukhin, I. (2017). Attention is all you need. Advances in Neural Information Processing Systems, 30.",
        "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., … & Kiela, D. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. Advances in Neural Information Processing Systems, 33, 9459-9474.",
        "Brown, T., Mann, B., Ryder, N., Subbiah, M., Kaplan, J. D., Dhariwal, P., … & Amodei, D. (2020). Language models are few-shot learners. Advances in Neural Information Processing Systems, 33, 1877-1901.",
        "Shahul Es, J. James, L. Espinosa-Anke, S. Schockaert (2023). RAGAS: Automated evaluation of Retrieval Augmented Generation. arXiv:2309.15217.",
        "LangChain Inc. (2024). LangChain documentation. Recuperado de https://python.langchain.com/docs/",
        "LangChain Inc. (2024). LangGraph documentation. Recuperado de https://langchain-ai.github.io/langgraph/",
        "Microsoft Corporation. (2024). Azure AI Search documentation. Recuperado de https://learn.microsoft.com/azure/search/",
        "OpenAI. (2024). Azure OpenAI Service documentation. Recuperado de https://learn.microsoft.com/azure/ai-services/openai/",
        "Kasneci, E., Seßler, K., Küchemann, S., Bannert, M., Dementieva, D., Fischer, F., … & Kasneci, G. (2023). ChatGPT for good? On opportunities and challenges of large language models for education. Learning and Individual Differences, 103, 102274.",
        "Grassini, S. (2023). Shaping the future of education: Exploring the potential and consequences of AI and ChatGPT in educational settings. Education Sciences, 13(7), 692.",
        "Wollny, S., Schneider, J., Di Mitri, D., Weidlich, J., Rittberger, M., & Drachsler, H. (2021). Are we there yet? A systematic literature review on chatbots in education. Frontiers in Artificial Intelligence, 4, 654924.",
        "Okonkwo, C. W., & Ade-Ibijola, A. (2021). Chatbots applications in education: A systematic review. Computers and Education: Artificial Intelligence, 2, 100033.",
        "Gao, Y., Xiong, Y., Gao, X., Jia, K., Pan, J., Bi, Y., … & Wang, H. (2024). Retrieval-augmented generation for large language models: A survey. arXiv:2312.10997.",
        "Karpukhin, V., Oguz, B., Min, S., Lewis, P., Wu, L., Edunov, S., … & Yih, W.-t. (2020). Dense passage retrieval for open-domain question answering. arXiv:2004.04906.",
        "Wei, J., Wang, X., Schuurmans, D., Bosma, M., Xia, F., Chi, E., … & Zhou, D. (2022). Chain-of-thought prompting elicits reasoning in large language models. Advances in Neural Information Processing Systems, 35, 24824-24837.",
        "Fan, W., Ding, Y., Ning, L., Wang, S., Li, H., Yin, D., … & Li, Q. (2024). A survey on RAG meeting LLMs: Towards retrieval-augmented large language models. arXiv:2405.06211.",
        "Wu, Q., Bansal, G., Zhang, J., Wu, Y., Li, B., Zhu, E., … & Wang, C. (2023). AutoGen: Enabling next-gen LLM applications via multi-agent conversation. arXiv:2308.08155.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(f"[{i}] {ref}")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)
        for run in p.runs:
            run.font.size = Pt(11)

    # ══════════════════════════════════════════════════════════════════════════
    # ANEXO A — ESTRUCTURA DEL PROYECTO
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Anexo A — Estructura del Proyecto", 1)
    add_body(doc, "Árbol de directorios del proyecto tal como existe en disco a fecha de la Entrega 1:")
    add_code_block(doc,
"E:\\TFG\\tfg-universidad-rag\\\n"
"├── .env.example              → Template de variables de entorno (sin claves reales)\n"
"├── .gitignore                → Excluye .env, *.db, logs/, __pycache__/, etc.\n"
"├── CLAUDE.md                 → Instrucciones operativas para Claude Code\n"
"├── README.md                 → Descripción, instalación y comandos\n"
"├── requirements.txt          → 13 dependencias Python con versiones mínimas\n"
"├── app/\n"
"│   ├── __init__.py\n"
"│   └── components/\n"
"│       └── __init__.py       → (Fase 4: chat.py, sidebar.py, sources_panel.py)\n"
"├── data/\n"
"│   ├── corpus/\n"
"│   │   ├── normativas/       → 3 documentos: reglamento, permanencia, evaluación\n"
"│   │   ├── planes_estudio/   → 3 documentos: Informática, ADE, Derecho\n"
"│   │   ├── guias_docentes/   → 8 fichas de asignaturas (INF101-INF304)\n"
"│   │   ├── faqs/             → 4 colecciones: matrícula, becas, admisión, exámenes\n"
"│   │   └── procedimientos/   → 4 documentos: orientación, riesgo, servicios, calendario\n"
"│   └── database/\n"
"│       └── students.db       → SQLite: 50 alumnos, 1114 asignaturas, 24 becas\n"
"├── docs/\n"
"│   └── diagrams/             → (Fase 3: diagramas Mermaid para la memoria)\n"
"├── logs/\n"
"│   └── fase1_20260423.txt    → Log de verificación y resumen de sesión\n"
"├── scripts/\n"
"│   ├── generate_students.py  → Genera y puebla students.db\n"
"│   └── verify_phase1.py      → 21 checks de verificación de Fase 1\n"
"├── src/\n"
"│   ├── agents/               → (Fase 3: router.py, guardrail.py, retriever.py...)\n"
"│   ├── config/\n"
"│   │   ├── azure_config.py   → Inicialización lazy de clientes Azure\n"
"│   │   └── settings.py       → Constantes globales del sistema\n"
"│   ├── graph/                → (Fase 3: state.py, graph.py, edges.py)\n"
"│   ├── prompts/              → (Fase 3: prompts de cada agente)\n"
"│   └── tools/\n"
"│       └── sqlite_query.py   → 8 funciones de consulta SQLite predefinidas\n"
"└── tests/                    → (Fase 5: test_agents.py, test_guardrails.py...)")

    # ══════════════════════════════════════════════════════════════════════════
    # ANEXO B — FRAGMENTOS DE CÓDIGO
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Anexo B — Fragmentos de Código Relevantes", 1)

    add_heading(doc, "B.1 src/config/settings.py (completo)", 2)
    add_code_block(doc,
'"""Constantes globales del sistema. Lee valores de variables de entorno con defaults seguros."""\n\n'
'import os\n'
'from dotenv import load_dotenv\n\n'
'load_dotenv()\n\n'
'MAX_RETRIEVED_DOCS: int = int(os.getenv("MAX_RETRIEVED_DOCS", "5"))\n'
'RETRIEVAL_SCORE_THRESHOLD: float = float(os.getenv("RETRIEVAL_SCORE_THRESHOLD", "0.7"))\n'
'CHUNK_SIZE: int = 512\n'
'CHUNK_OVERLAP: int = 50\n'
'EMBEDDING_MODEL: str = "text-embedding-3-large"\n'
'EMBEDDING_DIMENSIONS: int = 3072\n'
'INDEX_NAME: str = os.getenv("AZURE_SEARCH_INDEX_NAME", "university-corpus")\n'
'SQLITE_DB_PATH: str = os.getenv("SQLITE_DB_PATH", "data/database/students.db")\n'
'CONVERSATIONS_DB_PATH: str = "data/database/conversations.db"\n'
'LOG_DIR: str = "logs"\n'
'DIAGRAMS_DIR: str = "docs/diagrams"')

    add_heading(doc, "B.2 src/tools/sqlite_query.py — función get_full_student_record (completa)", 2)
    add_code_block(doc,
'def get_full_student_record(student_id: str) -> dict:\n'
'    """Consolida toda la información del alumno en un único dict estructurado.\n\n'
'    Este es el único método que llama el Agente Student Data al State.\n\n'
'    Args:\n'
'        student_id: Identificador del alumno (ej. \'ALU001\').\n\n'
'    Returns:\n'
'        Dict con claves: profile, academic_standing, passed_subjects,\n'
'        pending_subjects, current_enrollments, scholarships.\n'
'    """\n'
'    logger.debug("get_full_student_record(%s)", student_id)\n\n'
'    profile = get_student_profile(student_id)\n'
'    if "error" in profile:\n'
'        return {"error": profile["error"]}\n\n'
'    academic_standing = get_academic_standing(student_id)\n'
'    passed_subjects    = get_passed_subjects(student_id)\n'
'    pending_subjects   = get_pending_subjects(student_id)\n'
'    current_enrollments = get_current_enrollments(student_id)\n'
'    scholarships       = get_scholarships(student_id)\n\n'
'    return {\n'
'        "profile": {\n'
'            "name": profile.get("name"),\n'
'            "degree": profile.get("degree"),\n'
'            "year": profile.get("year"),\n'
'            "gpa": profile.get("gpa"),\n'
'            "status": profile.get("status"),\n'
'        },\n'
'        "academic_standing": {\n'
'            "credits_completed":    academic_standing.get("credits_completed", 0),\n'
'            "credits_total":        academic_standing.get("credits_total", 240),\n'
'            "progress_pct":         academic_standing.get("progress_pct", 0.0),\n'
'            "at_risk":              academic_standing.get("at_risk", False),\n'
'            "failed_subjects_count": academic_standing.get("failed_subjects_count", 0),\n'
'        },\n'
'        "passed_subjects":    [{"code": s["subject_code"], "name": s["subject_name"],\n'
'                               "grade": s["grade"]} for s in passed_subjects\n'
'                               if "error" not in s],\n'
'        "pending_subjects":   [{"code": s["subject_code"], "name": s["subject_name"],\n'
'                               "status": s["status"]} for s in pending_subjects\n'
'                               if "error" not in s],\n'
'        "current_enrollments":[{"code": s["subject_code"], "name": s["subject_name"],\n'
'                               "semester": s["semester"]} for s in current_enrollments\n'
'                               if "error" not in s],\n'
'        "scholarships":       [{"type": s["type"], "status": s["status"],\n'
'                               "amount": s["amount"]} for s in scholarships\n'
'                               if "error" not in s],\n'
'    }')

    add_heading(doc, "B.3 scripts/generate_students.py — esquema y alumnos especiales", 2)
    add_code_block(doc,
'def create_schema(conn: sqlite3.Connection) -> None:\n'
'    """Crea las tablas en la base de datos."""\n'
'    conn.executescript("""\n'
'        DROP TABLE IF EXISTS scholarships;\n'
'        DROP TABLE IF EXISTS subject_records;\n'
'        DROP TABLE IF EXISTS students;\n\n'
'        CREATE TABLE students (\n'
'            id                TEXT PRIMARY KEY,\n'
'            name              TEXT NOT NULL,\n'
'            email             TEXT UNIQUE,\n'
'            degree            TEXT,\n'
'            year              INTEGER,\n'
'            gpa               REAL,\n'
'            credits_completed INTEGER,\n'
'            credits_total     INTEGER,\n'
'            status            TEXT,\n'
'            enrolled_year     INTEGER\n'
'        );\n'
'        ...\n'
'    """)\n\n\n'
'# Alumnos especiales fijados por el plan de pruebas:\n'
'SPECIAL: dict[str, tuple[str, int, str]] = {\n'
'    "ALU001": ("Informatica", 4, "excellent"),  # GPA 9.1, matriculado en TFG\n'
'    "ALU015": ("Informatica", 2, "at_risk"),    # GPA 4.2, asignaturas suspensas\n'
'    "ALU030": ("ADE", 3, "active"),             # GPA 7.8, beca MEC 1500€\n'
'    "ALU045": ("Derecho", 1, "active"),         # GPA 5.5, sin historial\n'
'}')

    # ══════════════════════════════════════════════════════════════════════════
    # ANEXO C — CORPUS DOCUMENTAL
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Anexo C — Corpus Documental", 1)
    add_body(doc, "Catálogo completo de los documentos que componen la base de conocimiento universitaria sintética:")
    add_table(doc,
        ["Archivo", "Categoría", "Palabras", "Secciones principales"],
        [
            ["reglamento_academico.md", "normativa", "1.544", "Preámbulo, Arts. 1-12 (matrícula, convocatorias, evaluación, permanencia, TFG)"],
            ["normativa_permanencia.md", "normativa", "1.053", "Arts. 1-7 (progreso, baja académica, notificación, PMP, readmisión)"],
            ["politica_evaluacion.md", "normativa", "1.134", "Arts. 1-7 (calificación, evaluación continua, asistencia, publicación, revisión)"],
            ["grado_informatica.md", "plan_estudio", "1.235", "Plan 4 años, tabla de asignaturas por curso (40+), catálogo de optativas (19)"],
            ["grado_ade.md", "plan_estudio", "1.146", "Plan 4 años, tabla de asignaturas (35+), catálogo de optativas ADE (19)"],
            ["grado_derecho.md", "plan_estudio", "1.178", "Plan 4 años, tabla de asignaturas (35+), catálogo de optativas Derecho (19)"],
            ["guia_programacion_i.md", "guia_docente", "617", "Identificación, Objetivos, Temario 6T, Metodología, Evaluación, Bibliografía"],
            ["guia_estructuras_datos.md", "guia_docente", "606", "Identificación, Objetivos, Temario 6T, Metodología, Evaluación, Bibliografía"],
            ["guia_bases_datos.md", "guia_docente", "624", "Identificación, Objetivos, Temario 6T, Metodología, Evaluación, Bibliografía"],
            ["guia_redes.md", "guia_docente", "616", "Identificación, Objetivos, Temario 6T, Metodología, Evaluación, Bibliografía"],
            ["guia_ia_introduccion.md", "guia_docente", "624", "Identificación, Objetivos, Temario 6T, Metodología, Evaluación, Bibliografía"],
            ["guia_matematicas_i.md", "guia_docente", "626", "Identificación, Objetivos, Temario 6T, Metodología, Evaluación, Bibliografía"],
            ["guia_calculo.md", "guia_docente", "587", "Identificación, Objetivos, Temario 6T, Metodología, Evaluación, Bibliografía"],
            ["guia_estadistica.md", "guia_docente", "609", "Identificación, Objetivos, Temario 6T, Metodología, Evaluación, Bibliografía"],
            ["faq_matriculacion.md", "faq", "1.055", "15 pares P/R: plazos, anulación, cambio de grupo, créditos, condicional"],
            ["faq_becas.md", "faq", "1.115", "12 pares P/R: beca MEC (req. académicos, económicos), excelencia, Erasmus"],
            ["faq_admision.md", "faq", "967", "10 pares P/R: notas de corte (Inf. 8.2, ADE 7.5, Der. 7.0), proceso, FP"],
            ["faq_examenes.md", "faq", "1.042", "12 pares P/R: convocatorias, revisión, reclamación, compensación, septiembre"],
            ["guia_orientacion_academica.md", "procedimiento", "798", "Qué es el SOA, para qué sirve, cómo pedir cita, protocolo primera sesión"],
            ["protocolo_apoyo_riesgo.md", "procedimiento", "1.094", "Criterios detección, notificación, tutoría obligatoria, PMP, consecuencias"],
            ["servicios_universidad.md", "procedimiento", "1.270", "Psicología, Tutorías, Biblioteca, Deportes, Erasmus, Empleo, Secretaría"],
            ["calendario_academico.md", "procedimiento", "856", "Fechas 2024-25: matrícula, semestres, exámenes, TFG, becas"],
        ]
    )

    # Guardar
    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Documento guardado en: {OUT_PATH}")


if __name__ == "__main__":
    build_document()
